# utils.py
# -*- coding: utf-8 -*-

import re
from collections import Counter
from typing import List, Set, Tuple, Iterable

# =============================================================================
#               ЧТЕНИЕ ТЕКСТА ИЗ ФАЙЛОВ (PDF / DOCX / TXT)
# =============================================================================

async def extract_text_from_file(file) -> str:
    """
    Принимает FastAPI UploadFile объект и вытягивает текст из PDF/DOCX/TXT.
    Возвращает строку (может быть пустой, если парсинг не удался).
    """
    if not file:
        return ""

    name = (getattr(file, "filename", "") or "").lower()

    # --- TXT ---
    if name.endswith(".txt"):
        try:
            data = await file.read()
            return data.decode("utf-8", errors="ignore")
        except Exception as e:
            print(f"Error reading TXT file: {e}")
            return ""

    # --- PDF ---
    if name.endswith(".pdf"):
        try:
            from PyPDF2 import PdfReader  # pyright: ignore[reportMissingImports]
        except Exception:
            return ""
        try:
            # For FastAPI UploadFile, we need to seek to beginning
            file.seek(0)
            reader = PdfReader(file)
            parts: List[str] = []
            
            print(f"PDF has {len(reader.pages)} pages")
            for i, p in enumerate(reader.pages):
                try:
                    t = p.extract_text() or ""
                    print(f"Page {i+1}: extracted {len(t)} characters")
                    if t and t.strip():
                        parts.append(t.strip())
                except Exception as page_error:
                    print(f"Error extracting text from page {i+1}: {page_error}")
                    continue
            
            # If PyPDF2 failed, try alternative method
            if not parts:
                print("PyPDF2 failed, trying alternative PDF extraction...")
                file.seek(0)
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(stream=file, filetype="pdf")
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        text = page.get_text()
                        if text and text.strip():
                            parts.append(text.strip())
                            print(f"PyMuPDF Page {page_num+1}: extracted {len(text)} characters")
                    doc.close()
                except ImportError:
                    print("PyMuPDF not available for alternative PDF extraction")
                except Exception as alt_error:
                    print(f"Alternative PDF extraction failed: {alt_error}")
            
            result = "\n".join(parts)
            print(f"Total PDF text extracted: {len(result)} characters")
            
            # If still no text, provide helpful error message
            if not result.strip():
                print("WARNING: No text could be extracted from PDF")
                print("This might be an image-based PDF or have special formatting")
                return "IMAGE_BASED_PDF_DETECTED"
            
            return result
            
        except Exception as e:
            print(f"Error reading PDF file: {e}")
            return ""

    # --- DOCX ---
    if name.endswith(".docx"):
        try:
            import docx  # python-docx  # pyright: ignore[reportMissingImports]
            print("Successfully imported python-docx")
        except Exception as e:
            print(f"Failed to import python-docx: {e}")
            return ""
        try:
            # For FastAPI UploadFile, we need to read the content first
            file.seek(0)
            file_content = await file.read()
            print(f"DOCX file content read, size: {len(file_content)} bytes")
            
            # Create a BytesIO object for docx.Document
            from io import BytesIO
            file_stream = BytesIO(file_content)
            
            doc = docx.Document(file_stream)
            print(f"DOCX document opened, has {len(doc.paragraphs)} paragraphs")
            
            text_parts = []
            for i, p in enumerate(doc.paragraphs):
                if p.text and p.text.strip():
                    text_parts.append(p.text.strip())
                    print(f"Paragraph {i+1}: '{p.text[:50]}...'")
            
            result = "\n".join(text_parts)
            print(f"DOCX extraction result: {len(result)} characters")
            return result
            
        except Exception as e:
            print(f"Error reading DOCX file: {e}")
            return ""

    # --- Fallback: пробуем прочитать как текст ---
    try:
        data = await file.read()
        return data.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"Error in fallback file reading: {e}")
        return ""

def check_ats_readability(resume_text: str, file_name: str = "") -> dict:
    """
    Проверяет резюме на совместимость с ATS системами.
    Возвращает словарь с результатами проверки и рекомендациями.
    """
    if not resume_text:
        return {
            "score": 0,
            "status": "critical",
            "issues": ["No text could be extracted from the resume"],
            "recommendations": ["Check if the file format is supported", "Try converting to a different format"],
            "ats_friendly": False
        }
    
    # Calculate readability metrics
    total_chars = len(resume_text)
    total_words = len(resume_text.split())
    total_lines = len(resume_text.split('\n'))
    
    # Check for common ATS issues
    issues = []
    recommendations = []
    
    # Issue 1: Text too short (might be image-based PDF)
    if total_words < 50:
        issues.append("Very short text extracted - might be an image-based PDF")
        recommendations.append("Convert to text-based PDF or use Word document")
    
    # Issue 2: Too many special characters (might be OCR artifacts)
    special_char_ratio = len(re.findall(r'[^\w\s]', resume_text)) / max(1, total_chars)
    if special_char_ratio > 0.3:
        issues.append("High number of special characters - might be OCR artifacts")
        recommendations.append("Check if text was properly extracted")
    
    # Issue 3: Check for common ATS-unfriendly patterns (less strict)
    ats_unfriendly_patterns = [
        (r'[^\x00-\x7F]', "Non-ASCII characters detected"),
        (r'\b[A-Z]{5,}\b', "Excessive all-caps text detected (harder for ATS to parse)"),
        # Only flag really unusual characters, not common punctuation
        (r'[^\w\s\-\.\,\;\:\!\?\(\)\[\]\{\}\"\']', "Unusual formatting characters detected"),
    ]
    
    for pattern, issue_desc in ats_unfriendly_patterns:
        if re.search(pattern, resume_text):
            issues.append(issue_desc)
    
    # Issue 4: Check if text looks like it was properly extracted
    # Look for signs of image-based PDF or design tool export
    suspicious_patterns = [
        (r'^[^\w]*$', "Lines with no readable text"),
        (r'[A-Za-z]{25,}', "Very long words (might be design artifacts)"),
        (r'\n{5,}', "Excessive line breaks (design tool formatting)"),
    ]
    
    for pattern, issue_desc in suspicious_patterns:
        if re.search(pattern, resume_text):
            issues.append(issue_desc)
    
    # File format specific warnings
    if file_name and file_name.lower().endswith('.pdf'):
        if total_words < 100:
            issues.append("PDF appears to be image-based or design tool export")
            recommendations.append("Export as 'text-based PDF' from your design tool")
            recommendations.append("Or save as Word document (.docx) for better ATS compatibility")
    
    # Calculate overall ATS compatibility score
    score = 100
    
    # Deduct points for each issue
    if total_words < 100:
        score -= 30
    elif total_words < 200:
        score -= 15
    
    if special_char_ratio > 0.3:
        score -= 20
    elif special_char_ratio > 0.2:
        score -= 10
    
    score -= len(issues) * 10
    score = max(0, score)
    
    # Determine status
    if score >= 80:
        status = "excellent"
        ats_friendly = True
    elif score >= 60:
        status = "good"
        ats_friendly = True
    elif score >= 40:
        status = "fair"
        ats_friendly = False
    else:
        status = "poor"
        ats_friendly = False
    
    # Add general recommendations based on score
    if not ats_friendly:
        recommendations.extend([
            "Use simple, clean formatting",
            "Avoid complex layouts and graphics",
            "Use standard fonts (Arial, Calibri, Times New Roman)",
            "Save as .docx or text-based PDF",
            "Test with online ATS checkers"
        ])
    
    return {
        "score": score,
        "status": status,
        "ats_friendly": ats_friendly,
        "issues": issues,
        "recommendations": recommendations,
        "metrics": {
            "total_chars": total_chars,
            "total_words": total_words,
            "total_lines": total_lines,
            "special_char_ratio": round(special_char_ratio, 3)
        }
    }


# =============================================================================
#                           НОРМАЛИЗАЦИЯ / ТОКЕНЫ
# =============================================================================

def clean_text(s: str) -> str:
    """Грубая очистка: убрать \x00, схлопнуть пробелы и нормализовать переносы."""
    if not s:
        return ""
    s = s.replace("\x00", " ")
    

    
    # Then handle the rest normally
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\r\n?", "\n", s)
    return s.strip()


def detect_pdf_formatting_issues(resume_text: str) -> dict:
    """Detect common PDF formatting issues that break ATS systems."""
    issues = []
    recommendations = []
    severity = "low"
    
    # Check for the Figma/design tool export issue
    problematic_lines = []
    for i, line in enumerate(resume_text.split('\n'), 1):
        # Look for lines with many single letters separated by spaces
        if re.search(r'([a-zA-Z]\s+){3,}[a-zA-Z]', line):
            problematic_lines.append((i, line.strip()))
    
    if problematic_lines:
        severity = "high"
        issues.append(f"**PDF Export Issue Detected:** {len(problematic_lines)} lines have broken text formatting")
        issues.append("This commonly happens when exporting PDFs from design tools like Figma, Canva, or Photoshop")
        
        # Show examples of the problem
        issues.append("**Examples of broken text:**")
        for line_num, line in problematic_lines[:3]:  # Show first 3 examples
            issues.append(f"Line {line_num}: `{line[:50]}{'...' if len(line) > 50 else ''}`")
        
        recommendations.append("**Fix this in your design tool:**")
        recommendations.append("• Export as 'text-based PDF' instead of 'image-based PDF'")
        recommendations.append("• Save as Word document (.docx) for better ATS compatibility")
        recommendations.append("• Use 'Print to PDF' option if available")
        recommendations.append("• Check your design tool's PDF export settings")
        
        recommendations.append("**Why this happens:**")
        recommendations.append("Design tools often create PDFs where text is treated as graphics")
        recommendations.append("This makes the text unreadable by ATS systems and job platforms")
    
    return {
        "has_issues": len(problematic_lines) > 0,
        "severity": severity,
        "issues": issues,
        "recommendations": recommendations,
        "problematic_lines_count": len(problematic_lines)
    }


def create_ats_preview(resume_text: str, sections_found: Set[str]) -> dict:
    """Create an ATS preview showing how the document looks to ATS systems."""
    preview = {
        "sections": {},
        "content_sample": "",
        "ats_score": 0,
        "issues": [],
        "recommendations": []
    }
    
    # Analyze each section with better detection
    section_patterns = {
        'summary': r'(?i)(summary|objective|profile|overview)',
        'experience': r'(?i)(experience|work|employment|career|professional)',
        'skills': r'(?i)(skills|competencies|technologies|tools|expertise)',
        'education': r'(?i)(education|academic|degree|university|college)',
        'projects': r'(?i)(projects|portfolio|work|case studies|achievements)',
        'contact': r'(?i)(contact|email|phone|linkedin|address)'
    }
    
    lines = resume_text.split('\n')
    current_section = "header"
    section_content = []
    section_start_line = 0
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # Check if this line starts a new section
        section_found = False
        for section_name, pattern in section_patterns.items():
            if re.search(pattern, line):
                # Save previous section
                if current_section != "header" and section_content:
                    content_text = "\n".join(section_content)
                    word_count = len(content_text.split())
                    
                    # Better quality assessment
                    if word_count >= 10:
                        quality = "excellent"
                    elif word_count >= 5:
                        quality = "good"
                    elif word_count >= 2:
                        quality = "fair"
                    else:
                        quality = "poor"
                    
                    preview["sections"][current_section] = {
                        "content": content_text,
                        "found": True,
                        "quality": quality,
                        "word_count": word_count,
                        "start_line": section_start_line,
                        "end_line": i - 1,
                        "issues": []
                    }
                
                # Start new section
                current_section = section_name
                section_content = [line]
                section_start_line = i
                section_found = True
                break
        
        if not section_found:
            section_content.append(line)
    
    # Save last section
    if current_section != "header" and section_content:
        content_text = "\n".join(section_content)
        word_count = len(content_text.split())
        
        if word_count >= 10:
            quality = "excellent"
        elif word_count >= 5:
            quality = "good"
        elif word_count >= 2:
            quality = "fair"
        else:
            quality = "poor"
        
        preview["sections"][current_section] = {
            "content": content_text,
            "found": True,
            "quality": quality,
            "word_count": word_count,
            "start_line": section_start_line,
            "end_line": len(lines) - 1,
            "issues": []
        }
    
    # Analyze issues and provide specific recommendations
    for section_name, section_data in preview["sections"].items():
        if section_data["quality"] == "poor":
            if section_data["word_count"] == 0:
                section_data["issues"].append("Section is completely empty")
                preview["recommendations"].append(f"Add content to {section_name.title()} section")
            elif section_data["word_count"] == 1:
                section_data["issues"].append("Section has only 1 word - needs more detail")
                preview["recommendations"].append(f"Expand {section_name.title()} section with more details")
            else:
                section_data["issues"].append(f"Section has only {section_data['word_count']} words - too brief")
                preview["recommendations"].append(f"Add more content to {section_name.title()} section")
    
    # Calculate ATS score based on section quality
    total_sections = len(preview["sections"])
    if total_sections == 0:
        preview["ats_score"] = 0
    else:
        quality_scores = {
            "excellent": 100,
            "good": 80,
            "fair": 60,
            "poor": 20
        }
        total_score = sum(quality_scores.get(s["quality"], 0) for s in preview["sections"].values())
        preview["ats_score"] = total_score / total_sections
    
    # Add content sample
    preview["content_sample"] = resume_text[:500] + "..." if len(resume_text) > 500 else resume_text
    
    return preview

def tokenize(text: str) -> List[str]:
    """Простая токенизация: только букво-цифровые токены длиной >= 2."""
    if not text:
        return []
    return [t for t in re.findall(r"\b\w+\b", text.lower()) if len(t) >= 2]


# =============================================================================
#                СТОП-ЛИСТЫ: компания/география/филлера/«пух»
# =============================================================================

COMPANY_SUFFIXES = r"(?:inc|llc|gmbh|ltd|plc|corp|co|sa|ag|oy|kk|pte|pty|s\.?r\.?l\.?|s\.?a\.?s\.?|asa|ab|bv|nv)\b"

# Общее «служебное» — предлоги/союзы/частицы и т.п.
FILLER_STOP: Set[str] = set("""
including include includes about across within among between various multiple
customer customers client clients team teams company companies business businesses
join hiring hire role roles position opportunity opportunities responsibilities requirements
will would may can could must should via using use uses used
and or the a an for with of to in on at by from as is are was were be been being
your you we they our their this that these those it its into acrosses
""".split())

# «пушистые»/общие/водяные слова (добавлены те, что мешали в тестах)
FLUFF_STOP: Set[str] = {
    # Общие пустые
    "real","meaningful","believe","working","make","take","one","what","join","explore",
    "solutions","solution","opportunity","impact","value","role","responsibility",
    "responsibilities","requirements","offer","offers","needed","desired","preferred",
    "skills","skill","experience","experiences","knowledge","understanding","background",
    "ability","capable","strong","excellent","good","great","big",
    
    # Remove problematic words that don't represent skills
    "how","ll","but","through","enjoy","comfortable","countries","jette","overgaard",
    
    # «Вода»-глаголы
    "help","work","drive","support","develop","improve","ensure","deliver",
    "collaborate","collaboration","provide","including","build","building","contribute",

    # Функциональные общие
    "team","teams","environment","organization","organizational","project","projects",
    "business","customers","client","clients","stakeholders","company","companies",

    # Прилагательные без смысла
    "fast","innovative","new","next","future","current","global","local",
    "international","stronger","better","best","leading","unique","key","important",

    # «Везде и ни о чем»
    "platform","system","systems","process","processes","approach","methods","tools",
    "culture","focus","needs","goal","vision","mission","strategy","strategic",

    # Немного частых слов
    "more","high","low","many","across","around","different","various",

    # Конкретные, мешавшие в примерах
    "service","services","sales","offerings",
    
    # Overused/irrelevant words that don't add value
    "video","videos","people","everyone","anyone","someone","everybody","anybody",
    "lives","life","stories","story","sharing","share","shares","shared",
    "worth","exceptional","unlimited","reach","craft","crafting","content",
    "looking","talented","individual","individuals","who","shares","passion",
    "passionate","believe","believes","believed","mission","dreams","dream",
    "startup","startups","founders","founder","influence","influences","influenced",
    "tech","technology","technologies","get","gets","got","getting","done",
    "back","founded","burning","multistreaming","solution","solutions",
    "inspires","inspired","inspiring","worldwide","world","wide","follow",
    "following","follows","followed","small","highly","driven","focused",
    "lasting","lasting","impact","impacts","impacted","offer","offers","offered",
    "closely","closer","close","build","building","built","grow","growing","grown",
    "evolution","evolve","evolved","evolving","create","creates","created","creating",
    "creator","creators","creation","creative","creativity","influences","influence",
    "influencing","influential","need","needs","needed","job","jobs","work","works",
    "working","worked","environment","environments","structure","structured",
    "flat","company","companies","corporate","corporation","corporations",
    
    # Additional generic/irrelevant words that don't represent skills
    "has","have","had","having","become","becomes","became","becoming",
    "top","main","challenge","challenges","challenging","transform","transforms",
    "transformed","transforming","level","levels","setting","set","sets",
    "standards","standard","hands","hand","find","finds","finding","found",
    "elegant","elegance","attention","detail","details","overall","direction",
    "directions","during","other","others","another","each","every","all",
    "some","many","few","several","various","different","similar","same",
    "first","second","third","next","previous","current","recent","new",
    "old","young","big","small","large","tiny","huge","enormous",
    "important","significant","essential","critical","crucial","vital",
    "major","minor","primary","secondary","tertiary","main","sub",
    "core","central","peripheral","additional","extra","supplementary",
    "complementary","supporting","assisting","helping","aiding","facilitating",
    "enabling","empowering","strengthening","enhancing","improving","bettering",
    "advancing","progressing","developing","evolving","growing","expanding",
    "extending","broadening","widening","deepening","intensifying","strengthening",
    "reinforcing","consolidating","unifying","integrating","connecting","linking",
    "joining","combining","merging","blending","mixing","combining","uniting",
    "bringing","taking","making","doing","performing","executing","carrying",
    "conducting","leading","guiding","directing","managing","overseeing",
    "supervising","coordinating","organizing","arranging","structuring",
    "planning","strategizing","thinking","considering","evaluating","assessing",
    "analyzing","examining","investigating","exploring","researching","studying",
    "learning","understanding","comprehending","grasping","realizing","recognizing",
    "identifying","discovering","finding","locating","spotting","noticing",
    "seeing","observing","watching","monitoring","tracking","following",
    "pursuing","chasing","seeking","searching","looking","finding","discovering"
}

# География/локали
GEO_STOP: Set[str] = set("""
nordic nordics europe european cet cest oslo helsinki stockholm norway sweden finland
germany france uk britain england denmark iceland scandinavia
""".split())

# Add new stoplist for noise words
NOISE_STOP = {
    # Basic noise words
    'over', 'give', 'complete', 'control', 'move', 'act', 'person', 'part',
    'now', 'where', 're', 've', 'do', 'll', 's', 't', 'm', 'd', 'nt',
    'get', 'make', 'take', 'put', 'set', 'let', 'hit', 'cut', 'run', 'sit',
    'stand', 'walk', 'talk', 'think', 'feel', 'know', 'see', 'hear', 'say',
    'tell', 'ask', 'answer', 'call', 'find', 'keep', 'hold', 'bring', 'carry',
    'send', 'show', 'read', 'write', 'draw', 'build', 'create', 'design',
    'plan', 'work', 'play', 'eat', 'drink', 'sleep', 'wake', 'open', 'close',
    'start', 'stop', 'begin', 'end', 'come', 'go', 'leave', 'arrive', 'reach',
    
    # Team and descriptive words
    'portfolio', 'senior', 'vibrant', 'marathon', 'runners', 'musicians', 'bakers', 
    'bookworms', 'football', 'players', 'united', 'team', 'teams', 'group', 'groups',
    'crew', 'squad', 'squads', 'band', 'bands', 'orchestra', 'orchestras',
    'ensemble', 'ensembles', 'club', 'clubs', 'association', 'associations',
    'community', 'communities', 'network', 'networks', 'family', 'families',
    'brotherhood', 'sisterhood', 'fellowship', 'fellowships', 'partnership', 'partnerships',
    'collaboration', 'collaborations', 'alliance', 'alliances', 'coalition', 'coalitions',
    'federation', 'federations', 'league', 'leagues', 'division', 'divisions',
    'department', 'departments', 'unit', 'units', 'section', 'sections',
    'branch', 'branches', 'outlet', 'outlets', 'store', 'stores', 'shop', 'shops',
    'market', 'markets', 'boutique', 'boutiques', 'gallery', 'galleries',
    'studio', 'studios', 'workshop', 'workshops', 'lab', 'labs', 'laboratory', 'laboratories',
    'center', 'centers', 'centre', 'centres', 'hub', 'hubs', 'node', 'nodes',
    'station', 'stations', 'post', 'posts', 'base', 'bases', 'camp', 'camps',
    'headquarters', 'headquarter', 'office', 'offices', 'building', 'buildings',
    'facility', 'facilities', 'complex', 'complexes', 'campus', 'campuses',
    'premises', 'premise', 'venue', 'venues', 'location', 'locations',
    'site', 'sites', 'area', 'areas', 'zone', 'zones', 'region', 'regions',
    'district', 'districts', 'neighborhood', 'neighborhoods', 'quarter', 'quarters',
    
    # Job description irrelevant words
    'together', 'faster', 'ambitions', 'active', 'passionate', 'enthusiastic', 'motivated',
    'dedicated', 'committed', 'reliable', 'trustworthy', 'honest', 'integrity', 'loyal',
    'flexible', 'adaptable', 'versatile', 'dynamic', 'energetic', 'creative', 'innovative',
    'strategic', 'tactical', 'operational', 'functional', 'practical', 'theoretical',
    'academic', 'scholarly', 'intellectual', 'analytical', 'logical', 'systematic',
    'methodical', 'organized', 'structured', 'planned', 'scheduled', 'timely', 'efficient',
    'effective', 'productive', 'proactive', 'initiative', 'self-starter', 'independent',
    'collaborative', 'cooperative', 'supportive', 'helpful', 'friendly', 'approachable',
    'communicative', 'articulate', 'eloquent', 'persuasive', 'influential', 'charismatic',
    'leadership', 'management', 'supervision', 'mentoring', 'coaching', 'training',
    'development', 'growth', 'improvement', 'enhancement', 'optimization', 'maximization',
    'minimization', 'reduction', 'increase', 'decrease', 'maintenance', 'sustainability',
    'stability', 'consistency', 'reliability', 'durability', 'longevity', 'permanence',
    'temporary', 'permanent', 'full-time', 'part-time', 'contract', 'freelance',
    'remote', 'hybrid', 'onsite', 'offsite', 'travel', 'relocation', 'mobility',
    'flexibility', 'adaptability', 'versatility', 'diversity', 'inclusion', 'equity',
    'fairness', 'justice', 'ethics', 'morality', 'values', 'principles', 'standards',
    'quality', 'excellence', 'superiority', 'inferiority', 'mediocrity', 'average',
    'outstanding', 'exceptional', 'remarkable', 'notable', 'significant', 'important',
    'critical', 'essential', 'vital', 'crucial', 'key', 'major', 'minor', 'primary',
    'secondary', 'tertiary', 'quaternary', 'quinary', 'senary', 'septenary', 'octonary',
    'nonary', 'denary', 'duodenary', 'vigesimal', 'sexagesimal', 'centesimal',
    'millesimal', 'microscopic', 'macroscopic', 'infinitesimal', 'infinite', 'finite',
    'limited', 'unlimited', 'restricted', 'unrestricted', 'bounded', 'unbounded',
    'constrained', 'unconstrained', 'regulated', 'unregulated', 'controlled', 'uncontrolled',
    'managed', 'unmanaged', 'supervised', 'unsupervised', 'guided', 'unguided',
    'directed', 'undirected', 'oriented', 'disoriented', 'focused', 'unfocused',
    'concentrated', 'distributed', 'centralized', 'decentralized', 'localized', 'globalized',
    'specialized', 'generalized', 'standardized', 'customized', 'personalized', 'individualized',
    'collectivized', 'socialized', 'privatized', 'nationalized', 'internationalized',
    'globalized', 'localized', 'regionalized', 'continentalized', 'planetary', 'interplanetary',
    'interstellar', 'intergalactic', 'cosmic', 'universal', 'eternal', 'temporal',
    'spatial', 'dimensional', 'multidimensional', 'bidimensional', 'tridimensional',
    'quadridimensional', 'pentadimensional', 'hexadimensional', 'heptadimensional',
    'octadimensional', 'enneadimensional', 'decadimensional', 'undecadimensional',
    'duodecadimensional', 'tredecadimensional', 'quattuordecadimensional', 'quindecadimensional',
    'sedecadimensional', 'septendecadimensional', 'octodecadimensional', 'novemdecadimensional',
    'vigintadimensional', 'unvigintadimensional', 'duovigintadimensional', 'trevigintadimensional',
    'quattuorvigintadimensional', 'quinvigintadimensional', 'sexvigintadimensional',
    'septenvigintadimensional', 'octovigintadimensional', 'novemvigintadimensional',
    'trigintadimensional', 'untrigintadimensional', 'duotrigintadimensional', 'tretrigintadimensional',
    'quattuortrigintadimensional', 'quintrigintadimensional', 'sextrigintadimensional',
    'septentrigintadimensional', 'octotrigintadimensional', 'novemtrigintadimensional',
    'quadragintadimensional', 'unquadragintadimensional', 'duoquadragintadimensional',
    'trequadragintadimensional', 'quattuorquadragintadimensional', 'quinquadragintadimensional',
    'sexquadragintadimensional', 'septenquadragintadimensional', 'octoquadragintadimensional',
    'novemquadragintadimensional', 'quinquagintadimensional', 'unquinquagintadimensional',
    'duoquinquagintadimensional', 'trequinquagintadimensional', 'quattuorquinquagintadimensional',
    'quinquinquagintadimensional', 'sexquinquagintadimensional', 'septenquinquagintadimensional',
    'octoquinquagintadimensional', 'novemquinquagintadimensional', 'sexagintadimensional',
    'unsexagintadimensional', 'duosexagintadimensional', 'tresexagintadimensional',
    'quattuorsexagintadimensional', 'quinsexagintadimensional', 'sexsexagintadimensional',
    'septensexagintadimensional', 'octosexagintadimensional', 'novemsexagintadimensional',
    'septuagintadimensional', 'unseptuagintadimensional', 'duoseptuagintadimensional',
    'treseptuagintadimensional', 'quattuorseptuagintadimensional', 'quinseptuagintadimensional',
    'sexseptuagintadimensional', 'septenseptuagintadimensional', 'octoseptuagintadimensional',
    'novemseptuagintadimensional', 'octogintadimensional', 'unoctogintadimensional',
    'duooctogintadimensional', 'treoctogintadimensional', 'quattuoroctogintadimensional',
    'quinoctogintadimensional', 'sexoctogintadimensional', 'septenoctogintadimensional',
    'octooctogintadimensional', 'novemoctogintadimensional', 'nonagintadimensional',
    'unnonagintadimensional', 'duononagintadimensional', 'trenonagintadimensional',
    'quattuornonagintadimensional', 'quinnonagintadimensional', 'sexnonagintadimensional',
    'septennonagintadimensional', 'octononagintadimensional', 'novemnonagintadimensional',
    'centadimensional', 'uncentadimensional', 'duocentadimensional', 'trecentadimensional',
    'quattuorcentadimensional', 'quincentadimensional', 'sexcentadimensional',
    'septencentadimensional', 'octocentadimensional', 'novemcentadimensional',
    
    # Additional job description irrelevant words
    'required', 'reality', 'obsessed', 'manager', 'way', 'industry', 'also', 'canada', 
    'made', 'while', 'cities', 'need', 'want', 'must', 'should', 'could', 'would',
    'might', 'may', 'can', 'will', 'shall', 'do', 'does', 'did', 'done', 'doing',
    'have', 'has', 'had', 'having', 'be', 'am', 'is', 'are', 'was', 'were', 'been',
    'being', 'become', 'becomes', 'became', 'becoming', 'get', 'gets', 'got', 'getting',
    'gotten', 'go', 'goes', 'went', 'gone', 'going', 'come', 'comes', 'came', 'coming',
    'make', 'makes', 'made', 'making', 'take', 'takes', 'took', 'taking', 'taken',
    'give', 'gives', 'gave', 'giving', 'given', 'put', 'puts', 'putting', 'set',
    'sets', 'setting', 'let', 'lets', 'letting', 'hit', 'hits', 'hitting', 'cut',
    'cuts', 'cutting', 'run', 'runs', 'running', 'ran', 'sit', 'sits', 'sitting',
    'sat', 'stand', 'stands', 'standing', 'stood', 'walk', 'walks', 'walking', 'walked',
    'talk', 'talks', 'talking', 'talked', 'think', 'thinks', 'thinking', 'thought',
    'feel', 'feels', 'feeling', 'felt', 'know', 'knows', 'knowing', 'knew', 'known',
    'see', 'sees', 'seeing', 'saw', 'seen', 'hear', 'hears', 'hearing', 'heard',
    'say', 'says', 'saying', 'said', 'tell', 'tells', 'telling', 'told',
    'ask', 'asks', 'asking', 'asked', 'answer', 'answers', 'answering', 'answered',
    'call', 'calls', 'calling', 'called', 'find', 'finds', 'finding', 'found',
    'keep', 'keeps', 'keeping', 'kept', 'hold', 'holds', 'holding', 'held',
    'bring', 'brings', 'bringing', 'brought', 'carry', 'carries', 'carrying', 'carried',
    'send', 'sends', 'sending', 'sent', 'show', 'shows', 'showing', 'showed', 'shown',
    'read', 'reads', 'reading', 'write', 'writes', 'writing', 'wrote', 'written',
    'draw', 'draws', 'drawing', 'drew', 'drawn', 'build', 'builds', 'building', 'built',
    'create', 'creates', 'creating', 'created', 'design', 'designs', 'designing', 'designed',
    'plan', 'plans', 'planning', 'planned', 'work', 'works', 'working', 'worked',
    'play', 'plays', 'playing', 'played', 'eat', 'eats', 'eating', 'ate', 'eaten',
    'drink', 'drinks', 'drinking', 'drank', 'drunk', 'sleep', 'sleeps', 'sleeping', 'slept',
    'wake', 'wakes', 'waking', 'woke', 'woken', 'open', 'opens', 'opening', 'opened',
    'close', 'closes', 'closing', 'closed', 'start', 'starts', 'starting', 'started',
    'stop', 'stops', 'stopping', 'stopped', 'begin', 'begins', 'beginning', 'began', 'begun',
    'end', 'ends', 'ending', 'ended', 'leave', 'leaves', 'leaving', 'left',
    'arrive', 'arrives', 'arriving', 'arrived', 'reach', 'reaches', 'reaching', 'reached'
}

# Hard hint keywords that get score boosts
HARD_HINT_KEYWORDS = {
    'figma', 'prototyping', 'accessibility', 'procurement', 'design system',
    'user research', 'usability testing', 'wireframing', 'sketch', 'invision',
    'principle', 'framer', 'protopie', 'marvel', 'balsamiq', 'adobe',
    'photoshop', 'illustrator', 'react', 'frontend', 'backend', 'api',
    'microservices', 'docker', 'kubernetes', 'aws', 'azure', 'gcp',
    'agile', 'scrum', 'kanban', 'lean', 'six sigma', 'design thinking'
}

def extract_keywords_with_scores(jd_text: str, top_n: int = 30) -> List[Tuple[str, float]]:
    """Extract keywords with scores based on frequency + bonuses, with similar word deduplication."""
    # Clean and tokenize
    cleaned = clean_text(jd_text)
    tokens = tokenize(cleaned)
    
    # Filter out stop words and noise
    filtered_tokens = []
    for token in tokens:
        token_lower = token.lower()
        if (token_lower not in FILLER_STOP and 
            token_lower not in FLUFF_STOP and 
            token_lower not in GEO_STOP and
            token_lower not in NOISE_STOP and
            len(token_lower) >= 3):  # Minimum 3 characters
            filtered_tokens.append(token_lower)
    
    # Count frequencies
    token_counts = Counter(filtered_tokens)
    
    # Detect company names to filter out
    companies = detect_company_names(jd_text)
    company_parts = set()
    for company in companies:
        company_parts.add(company.lower())
        company_parts.update(company.lower().split())
    
    # Score keywords and deduplicate similar words
    keyword_scores = []
    seen_variants = set()
    
    for token, count in token_counts.most_common(top_n * 3):  # Get more candidates for deduplication
        if token in company_parts:
            continue  # Skip company-related terms
        
        # Check if this is a variant of an already seen word
        is_variant = False
        for seen_word in seen_variants:
            if (token == seen_word + 's' or  # user -> users
                token == seen_word + 'ing' or  # design -> designing
                token == seen_word + 'ed' or   # design -> designed
                token == seen_word + 'er' or   # design -> designer
                token == seen_word + 'al' or   # design -> designal
                token == seen_word[:-1] or     # users -> user
                token == seen_word[:-2] or     # designing -> design
                token == seen_word[:-3] or     # designed -> design
                token == seen_word[:-4] or     # designer -> design
                token == seen_word[:-5]):      # designal -> design
                is_variant = True
                break
        
        if is_variant:
            continue  # Skip variants
        
        # Add to seen variants
        seen_variants.add(token)
        
        score = count  # Base score is frequency
        
        # Add section bonuses
        if any(section in token for section in ['design', 'ux', 'ui', 'user']):
            score += 2
        if any(section in token for section in ['system', 'process', 'workflow']):
            score += 1.5
        if any(section in token for section in ['research', 'testing', 'analysis']):
            score += 1.5
        if any(section in token for section in ['prototype', 'wireframe', 'mockup']):
            score += 2
        
        # Add hard hint boosts
        if token in HARD_HINT_KEYWORDS:
            score += 3
        
        # Add length bonus for meaningful terms
        if len(token) >= 6:
            score += 0.5
        
        # Add relevance bonus based on JD context
        if 'design' in jd_text.lower() and token in ['design', 'ux', 'ui', 'visual', 'interaction']:
            score += 1
        if 'development' in jd_text.lower() and token in ['development', 'coding', 'programming', 'software']:
            score += 1
        if 'research' in jd_text.lower() and token in ['research', 'analysis', 'testing', 'user']:
            score += 1
        if 'management' in jd_text.lower() and token in ['management', 'leadership', 'strategy', 'planning']:
            score += 1
        
        keyword_scores.append((token, score))
    
    # Sort by score and return top_n
    keyword_scores.sort(key=lambda x: x[1], reverse=True)
    return keyword_scores[:top_n]

def select_most_relevant_keywords(ranked_kw: List[Tuple[str, float]], jd_text: str, top_k: int = 15) -> List[str]:
    """Select the most relevant keywords for the specific JD, avoiding duplicates and focusing on relevance."""
    
    # Extract key themes from JD
    jd_lower = jd_text.lower()
    jd_themes = {
        'design': ['design', 'ux', 'ui', 'visual', 'interaction', 'prototype', 'wireframe', 'sketch', 'figma'],
        'development': ['development', 'coding', 'programming', 'software', 'frontend', 'backend', 'api', 'react'],
        'research': ['research', 'analysis', 'testing', 'user', 'usability', 'feedback', 'insights'],
        'management': ['management', 'leadership', 'strategy', 'planning', 'coordination', 'collaboration'],
        'business': ['business', 'product', 'market', 'customer', 'stakeholder', 'requirements']
    }
    
    # Score keywords by relevance to JD themes
    relevant_keywords = []
    seen_roots = set()
    
    for keyword, score in ranked_kw:
        if len(relevant_keywords) >= top_k:
            break
            
        keyword_lower = keyword.lower()
        
        # Check if this keyword is a root form we haven't seen
        is_new_root = True
        for seen_root in seen_roots:
            if (keyword_lower == seen_root or
                keyword_lower.startswith(seen_root) or
                seen_root.startswith(keyword_lower) or
                len(set(keyword_lower) & set(seen_root)) / max(len(keyword_lower), len(seen_root)) > 0.8):
                is_new_root = False
                break
        
        if not is_new_root:
            continue
        
        # Calculate relevance score
        relevance_score = score
        
        # Boost relevance based on JD themes
        for theme, theme_keywords in jd_themes.items():
            if any(theme_kw in keyword_lower for theme_kw in theme_keywords):
                relevance_score += 2
                break
        
        # Boost if keyword appears in JD multiple times
        keyword_count = jd_lower.count(keyword_lower)
        if keyword_count > 1:
            relevance_score += keyword_count * 0.5
        
        relevant_keywords.append((keyword, relevance_score))
        seen_roots.add(keyword_lower)
    
    # Sort by relevance score and return top keywords
    relevant_keywords.sort(key=lambda x: x[1], reverse=True)
    return [kw for kw, score in relevant_keywords[:top_k]]

def get_keyword_coverage_explanation(present_count: int, total_count: int, similarity_score: float) -> str:
    """Generate detailed explanation of how the score is calculated."""
    
    coverage_percent = (present_count / total_count * 100) if total_count > 0 else 0
    
    explanation = f"""
**📊 How Your Score is Calculated:**

**Keyword Coverage (70% of total score):** {coverage_percent:.1f}%
- **Found:** {present_count} out of {total_count} top JD keywords
- **Missing:** {total_count - present_count} keywords to add
- **Formula:** ({present_count} ÷ {total_count}) × 70 = {coverage_percent * 0.7:.1f} points

**Text Similarity (30% of total score):** {similarity_score * 100:.1f}%
- **TF-IDF cosine similarity** between JD and resume texts
- **Formula:** {similarity_score * 100:.1f} × 0.3 = {similarity_score * 30:.1f} points

**Final Score:** {coverage_percent * 0.7 + similarity_score * 30:.1f} + {similarity_score * 30:.1f} = **{coverage_percent * 0.7 + similarity_score * 30:.0f}/100**

**💡 To improve your score:**
- **Add missing keywords** to increase coverage from {coverage_percent:.1f}% to 100%
- **Enhance resume content** to improve text similarity
- **Use specific examples** and metrics in your descriptions
"""
    
    return explanation

def smart_bullets_for_missing(missing: List[str]) -> List[str]:
    """Generate human-like example bullets for missing keywords."""
    
    # Specific keyword mappings
    keyword_templates = {
        'visual': [
            "• Established **visual design** standards that improved brand consistency across 5+ products",
            "• Created **visual** assets and design systems used by 20+ designers",
            "• Led **visual** design reviews that improved design quality by 40%"
        ],
        'consistency': [
            "• Implemented **consistency** guidelines that reduced design iterations by 30%",
            "• Built **consistency** frameworks that improved team efficiency by 25%",
            "• Maintained **consistency** across multiple platforms and devices"
        ],
        'figma': [
            "• Led **Figma** adoption across design team, improving collaboration by 50%",
            "• Created **Figma** component libraries used by 15+ designers",
            "• Established **Figma** workflows that reduced design handoff time by 60%"
        ],
        'ideation': [
            "• Facilitated **ideation** workshops with 20+ stakeholders",
            "• Led **ideation** sessions that generated 50+ innovative solutions",
            "• Applied **ideation** techniques to solve complex user experience challenges"
        ],
        'insight': [
            "• Conducted user research that provided **insights** driving 3 major product decisions",
            "• Generated **insights** from analytics data that improved conversion by 35%",
            "• Shared **insights** with stakeholders that influenced product roadmap"
        ],
        'research': [
            "• Conducted **research** with 100+ users to inform design decisions",
            "• Led **research** initiatives that identified 5 key user pain points",
            "• Applied **research** findings to improve user satisfaction scores by 45%"
        ],
        'prototyping': [
            "• Built interactive **prototypes** that accelerated stakeholder approval by 2 weeks",
            "• Created **prototyping** workflows that improved iteration speed by 40%",
            "• Used **prototyping** to test 10+ user flows before development"
        ],
        'workshops': [
            "• Facilitated **workshops** with cross-functional teams of 15+ people",
            "• Led **workshops** that generated actionable insights for 3 product features",
            "• Organized **workshops** that improved team alignment and collaboration"
        ],
        'design system': [
            "• Built **design system** components used across 8+ products",
            "• Established **design system** guidelines that improved consistency by 60%",
            "• Maintained **design system** that reduced design time by 30%"
        ],
        'frontend': [
            "• Collaborated with **frontend** developers to implement design solutions",
            "• Provided **frontend** specifications that improved development efficiency",
            "• Worked closely with **frontend** team to ensure design fidelity"
        ],
        'react': [
            "• Designed **React** components that improved development speed by 40%",
            "• Created **React**-compatible design specifications",
            "• Collaborated with **React** developers on component architecture"
        ],
        'streaming': [
            "• Designed **streaming** platform interfaces used by 100K+ users",
            "• Improved **streaming** user experience, reducing buffering complaints by 70%",
            "• Led **streaming** service redesign that increased engagement by 45%"
        ],
        'procurement': [
            "• Streamlined **procurement** processes that reduced costs by 25%",
            "• Implemented **procurement** systems that improved efficiency by 40%",
            "• Led **procurement** initiatives that saved $500K annually"
        ],
        'accessibility': [
            "• Implemented **accessibility** improvements that increased usability for 15% of users",
            "• Led **accessibility** audits that identified 20+ improvement opportunities",
            "• Established **accessibility** guidelines that ensured compliance with WCAG 2.1"
        ],
        'interaction': [
            "• Designed micro-**interactions** that improved user engagement by 35%",
            "• Created **interaction** patterns used across 10+ products",
            "• Defined **interaction** guidelines that improved consistency by 50%"
        ],
        'system': [
            "• Built **system** architectures that improved scalability by 300%",
            "• Established **system** guidelines that reduced errors by 45%",
            "• Maintained **system** components used by 50+ developers"
        ],
        'process': [
            "• Streamlined **process** workflows that improved efficiency by 40%",
            "• Implemented **process** improvements that reduced time-to-market by 3 weeks",
            "• Led **process** optimization initiatives that saved 200+ hours monthly"
        ],
        'workflow': [
            "• Designed **workflow** improvements that reduced task completion time by 50%",
            "• Optimized **workflow** processes that improved team productivity by 35%",
            "• Implemented **workflow** automation that eliminated 20+ manual steps"
        ]
    }
    
    bullets = []
    for keyword in missing:
        keyword_lower = keyword.lower()
        
        # Check for exact matches first
        if keyword_lower in keyword_templates:
            bullets.extend(keyword_templates[keyword_lower][:2])  # Take up to 2 bullets per keyword
        else:
            # Check for partial matches
            matched = False
            for template_key, template_bullets in keyword_templates.items():
                if template_key in keyword_lower or keyword_lower in template_key:
                    bullets.extend(template_bullets[:1])  # Take 1 bullet for partial matches
                    matched = True
                    break
            
            # Fallback for unknown keywords
            if not matched:
                bullets.append(f"• Incorporated **{keyword}** into project scope and delivery with measurable outcomes")
                bullets.append(f"• Applied **{keyword}** principles to improve team efficiency and project success")
    
    # Deduplicate and limit
    seen = set()
    unique_bullets = []
    for bullet in bullets:
        if bullet not in seen and len(unique_bullets) < 5:  # Limit to 5 bullets
            unique_bullets.append(bullet)
            seen.add(bullet)
    
    return unique_bullets

def _norm_company_name(s: str) -> str:
    s = s.lower()
    s = re.sub(r"[^\w\s.-]", " ", s)
    s = re.sub(rf"\b{COMPANY_SUFFIXES}\.?$", "", s.strip())
    return re.sub(r"[\s._-]+", " ", s).strip()

def detect_company_names(jd_text: str) -> Set[str]:
    """
    Пытаемся вытащить название компании из JD без внешних NLP-либ:
    - первые 3 строки заголовка
    - паттерны 'at X', 'About X', 'Join X', 'We at X'
    - паттерн 'X is|are hiring|looking'
    - домены в ссылках и email
    Возвращаем нормализованные имена.
    """
    cand: Set[str] = set()
    text = jd_text or ""
    lower = text.lower()

    # 1) первые три строки — там часто бренд
    head = "\n".join(text.splitlines()[:3])
    caps = re.findall(
        r"\b([A-Z][A-Za-z0-9&.-]*(?:\s+[A-Z][A-Za-z0-9&.-]*){0,2})(?:\s+" + COMPANY_SUFFIXES + r")?\b",
        head
    )
    for s in caps:
        cand.add(_norm_company_name(s))

    # 2) 'at X' / 'About X' / 'Join X' / 'We at X'
    for m in re.findall(r"\b(?:at|about|join|we at)\s+([A-Z][A-Za-z0-9&.-]*(?:\s+[A-Z][A-Za-z0-9&.-]*){0,2})", text):
        cand.add(_norm_company_name(m))

    # 3) 'X is|are hiring|looking'
    for m in re.findall(r"\b([A-Z][A-Za-z0-9&.-]*(?:\s+[A-Z][A-Za-z0-9&.-]*){0,2})\s+(?:is|are)\s+(?:hiring|looking)", text):
        cand.add(_norm_company_name(m))

    # 4) домены
    for dom in re.findall(r"https?://(?:www\.)?([a-z0-9-]+)\.(?:com|io|ai|co|tech|net|org|no|se|fi|uk|de|fr)\b", lower):
        cand.add(_norm_company_name(dom))
    for dom in re.findall(r"\b[a-z0-9._%+-]+@([a-z0-9-]+)\.(?:com|io|ai|co|tech|net|org|no|se|fi|uk|de|fr)\b", lower):
        cand.add(_norm_company_name(dom))

    # явный мусор
    cand -= {"careers", "jobs", "hiring", "company"}
    
    # Additional company-related terms to filter out
    company_related = {"restream", "restreaming", "multistream", "multistreaming"}
    cand -= company_related
    
    return {c for c in cand if c}


# =============================================================================
#                                КЛЮЧЕВЫЕ СЛОВА
# =============================================================================

def detect_role_type(jd_text: str) -> str:
    """
    Определяет тип роли на основе ключевых слов в JD.
    Возвращает строку с типом роли.
    """
    text_lower = jd_text.lower()
    
    # Product Design / UX/UI roles
    design_keywords = {"design", "designer", "ux", "ui", "user experience", "user interface", 
                       "prototype", "wireframe", "mockup", "design system", "visual design"}
    if any(keyword in text_lower for keyword in design_keywords):
        return "product_design"
    
    # Engineering roles
    eng_keywords = {"engineer", "engineering", "developer", "development", "programming", 
                    "code", "software", "technical", "architecture"}
    if any(keyword in text_lower for keyword in eng_keywords):
        return "engineering"
    
    # Product Management roles
    pm_keywords = {"product manager", "product management", "strategy", "roadmap", 
                   "requirements", "stakeholder", "business"}
    if any(keyword in text_lower for keyword in pm_keywords):
        return "product_management"
    
    # Marketing roles
    marketing_keywords = {"marketing", "growth", "acquisition", "campaign", "brand", 
                         "content", "social media", "analytics"}
    if any(keyword in text_lower for keyword in marketing_keywords):
        return "marketing"
    
    return "general"

def top_keywords(text: str, top_n: int = 30) -> List[str]:
    """Extract top keywords from text, with improved filtering."""
    # Clean and tokenize
    cleaned = clean_text(text)
    tokens = tokenize(cleaned)
    
    # Filter out stop words and noise
    filtered_tokens = []
    for token in tokens:
        token_lower = token.lower()
        if (token_lower not in FILLER_STOP and 
            token_lower not in FLUFF_STOP and 
            token_lower not in GEO_STOP and
            token_lower not in NOISE_STOP and
            len(token_lower) >= 3):  # Minimum 3 characters
            filtered_tokens.append(token_lower)
    
    # Count frequencies
    token_counts = Counter(filtered_tokens)
    
    # Get top keywords, ensuring no duplicates
    seen = set()
    top_keywords = []
    for token, count in token_counts.most_common():
        if token not in seen and len(top_keywords) < top_n:
            top_keywords.append(token)
            seen.add(token)
    
    return top_keywords


# =============================================================================
#                                МЕТРИКИ/СОВПАДЕНИЯ
# =============================================================================

def compute_keyword_overlap(resume_tokens: List[str], jd_keywords: List[str]) -> Tuple[int, int, float]:
    """
    Простой показатель: |∩|, |JD|, процент покрытия JD.
    """
    rset = set(resume_tokens)
    jset = set(jd_keywords)
    inter = len(rset & jset)
    total = max(1, len(jset))
    score = round(100.0 * inter / total, 1)
    return inter, total, score

def compute_similarity(resume_text: str, jd_text: str) -> float:
    """
    Вычисляет косинусное сходство между текстами резюме и JD (0..1).
    Использует простой TF-IDF подход без внешних библиотек.
    """
    if not resume_text or not jd_text:
        return 0.0
    
    # Токенизируем тексты
    resume_tokens = tokenize(resume_text)
    jd_tokens = tokenize(jd_text)
    
    if not resume_tokens or not jd_tokens:
        return 0.0
    
    # Создаем словари частот
    resume_freq = Counter(resume_tokens)
    jd_freq = Counter(jd_tokens)
    
    # Получаем все уникальные токены
    all_tokens = set(resume_freq.keys()) | set(jd_freq.keys())
    
    if not all_tokens:
        return 0.0
    
    # Вычисляем TF-IDF векторы (упрощенная версия)
    resume_vector = []
    jd_vector = []
    
    for token in all_tokens:
        # TF (term frequency) - частота токена в документе
        resume_tf = resume_freq.get(token, 0) / max(1, len(resume_tokens))
        jd_tf = jd_freq.get(token, 0) / max(1, len(jd_tokens))
        
        resume_vector.append(resume_tf)
        jd_vector.append(jd_tf)
    
    # Вычисляем косинусное сходство
    dot_product = sum(a * b for a, b in zip(resume_vector, jd_vector))
    
    resume_magnitude = sum(a * a for a in resume_vector) ** 0.5
    jd_magnitude = sum(a * a for a in jd_vector) ** 0.5
    
    if resume_magnitude == 0 or jd_magnitude == 0:
        return 0.0
    
    cosine_similarity = dot_product / (resume_magnitude * jd_magnitude)
    
    # Add safety checks to prevent extreme values
    if cosine_similarity > 1.0:
        # This can happen due to floating point precision issues
        cosine_similarity = 1.0
    elif cosine_similarity < 0.0:
        cosine_similarity = 0.0
    
    return cosine_similarity

def compute_keyword_similarity(resume_kw: Iterable[str], jd_kw: Iterable[str]) -> float:
    """
    Комбинированная метрика схожести (0..100):
    0.6 * покрытие JD  +  0.4 * Jaccard.
    """
    s1 = {str(x).lower().strip() for x in resume_kw if str(x).strip()}
    s2 = {str(x).lower().strip() for x in jd_kw if str(x).strip()}
    if not s1 or not s2:
        return 0.0

    inter = len(s1 & s2)
    union = len(s1 | s2)
    jaccard = inter / union if union else 0.0
    coverage = inter / len(s2) if s2 else 0.0

    score = 0.6 * coverage + 0.4 * jaccard
    return round(score * 100.0, 1)


# =============================================================================
#                       ПОДСКАЗКИ ПО ОТСУТСТВУЮЩИМ КЛЮЧАМ
# =============================================================================

def suggest_missing_keywords(
    jd_text: str,
    resume_text: str,
    top_n: int = 30,
    visibility_threshold: int = 1,
) -> Tuple[List[str], List[str], float]:
    """
    Возвращает кортеж (present, missing, coverage):
      • present — ключевые слова из JD, которые «видимы» в резюме (встречаются >= threshold),
      • missing — ключевые слова из JD, которых мало/нет в резюме,
      • coverage — интегральная оценка совпадения (0..100).
    """
    jd_keys: List[str] = top_keywords(jd_text, top_n=top_n)

    res_tokens = tokenize(resume_text)
    res_freq = Counter(res_tokens)
    
    # More flexible matching - check if keywords appear in resume text (case-insensitive)
    resume_text_lower = resume_text.lower()
    
    present = []
    missing = []
    
    # Debug: Let's see what we're actually working with
    print(f"DEBUG: JD keys to search for: {jd_keys}")
    print(f"DEBUG: Resume text length: {len(resume_text)}")
    print(f"DEBUG: Resume text sample (first 500 chars): {resume_text[:500]}")
    
    for keyword in jd_keys:
        # Check if keyword appears in resume text (more flexible than exact token match)
        keyword_lower = keyword.lower()
        if keyword_lower in resume_text_lower:
            present.append(keyword)
            print(f"DEBUG: FOUND '{keyword}' in resume")
        else:
            missing.append(keyword)
            print(f"DEBUG: MISSING '{keyword}' - not found in resume")
    
    print(f"DEBUG: Final results - Present: {present}, Missing: {missing}")

    coverage = compute_keyword_similarity(resume_kw=res_tokens, jd_kw=jd_keys)
    return present, missing, coverage


# =============================================================================
#                           ОБНАРУЖЕНИЕ СЕКЦИЙ РЕЗЮМЕ
# =============================================================================

def detect_sections(resume_text: str) -> Set[str]:
    """
    Обнаруживает стандартные секции резюме на основе заголовков.
    Возвращает множество найденных секций в нижнем регистре.
    """
    if not resume_text:
        return set()
    
    # Comprehensive section patterns with variations
    section_patterns = {
        'experience': [
            r'\b(?:experience|work\s+experience|employment|work\s+history|professional\s+experience|'
            r'career|work|employment\s+history|professional\s+background|work\s+background)\b'
        ],
        'skills': [
            r'\b(?:skills|technical\s+skills|competencies|expertise|capabilities|proficiencies|'
            r'technologies|tools|software|programming\s+languages|frameworks|libraries)\b'
        ],
        'education': [
            r'\b(?:education|academic|qualifications|degree|university|college|school|'
            r'certification|certifications|training|courses|learning|academic\s+background)\b'
        ],
        'projects': [
            r'\b(?:projects|portfolio|achievements|accomplishments|key\s+projects|'
            r'notable\s+projects|featured\s+work|case\s+studies|work\s+samples)\b'
        ],
        'summary': [
            r'\b(?:summary|profile|objective|about|overview|introduction|'
            r'professional\s+summary|career\s+objective|personal\s+statement)\b'
        ],
        'contact': [
            r'\b(?:contact|personal|links|portfolio|github|linkedin|email|phone|'
            r'address|location|social\s+media|online\s+presence|websites)\b'
        ],
        'languages': [
            r'\b(?:languages|certifications|courses|training|certificates|'
            r'licenses|accreditations|professional\s+development|continuing\s+education)\b'
        ],
        'interests': [
            r'\b(?:volunteer|interests|hobbies|activities|volunteer\s+work|'
            r'community\s+service|extracurricular|personal\s+interests|additional\s+activities)\b'
        ]
    }
    
    text_lower = resume_text.lower()
    lines = text_lower.split('\n')
    
    found_sections = set()
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # More flexible header detection - look for various formats
        # 1. Lines that start with capital letters
        # 2. Lines in ALL CAPS
        # 3. Lines with common section indicators
        # 4. Lines containing section keywords anywhere
        is_header = False
        
        # Check if line starts with capital letter
        if line and line[0].isupper():
            is_header = True
        # Check if line is in ALL CAPS (common in resumes)
        elif line.isupper() and len(line) > 2:
            is_header = True
        # Check if line has common section indicators
        elif any(indicator in line for indicator in ['experience', 'skills', 'education', 'projects', 'summary', 'contact']):
            is_header = True
        # Check if line contains section keywords anywhere (more flexible)
        elif any(keyword in line for keyword in ['experience', 'skills', 'education', 'projects', 'summary', 'contact', 'work', 'employment', 'career', 'background', 'portfolio', 'achievements']):
            is_header = True
        
        if is_header:
            # Check each section type
            for section_name, patterns in section_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, line):
                        found_sections.add(section_name)
                        break
                if section_name in found_sections:
                    break
    
    # Also check for sections that might be mentioned in the text content
    # This helps catch sections that might not have clear headers
    for section_name, patterns in section_patterns.items():
        if section_name not in found_sections:
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    found_sections.add(section_name)
                    break
    
    return found_sections
